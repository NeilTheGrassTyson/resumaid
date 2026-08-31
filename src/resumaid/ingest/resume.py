"""Resume ingestion: file -> text -> structured profile.

Two passes. Deterministic first — contact block, section headers, date ranges, degrees,
employer/title/date triples. An LLM pass handles only what the deterministic parse could not
segment, and every field it returns carries a pointer to its source span (ADR 0006).

Nothing is invented. A field with no source span is dropped, not guessed.

In the MVP resumes are *selected among*, never rewritten. This module extracts each document's
emphasis — what that variant leads with — so the matcher can name the best-fitting one per role.
"""

from __future__ import annotations

import hashlib
import re
import sqlite3
from pathlib import Path

from resumaid.models import Education, Employment, Profile, ResumeDoc
from resumaid.util import iso, jdump, jload, utcnow

SUPPORTED = {".pdf", ".docx", ".md", ".txt"}

_DEGREE_LEVELS = [
    ("doctorate", r"\b(ph\.?d|doctorate|doctoral|d\.?phil)\b"),
    ("masters", r"\b(m\.?s\.?c?|master'?s?|m\.?eng|m\.?b\.?a|m\.?a)\b"),
    ("bachelors", r"\b(b\.?s\.?c?|bachelor'?s?|b\.?eng|b\.?a|b\.?tech)\b"),
    ("associate", r"\b(a\.?a\.?s?|associate'?s? degree)\b"),
    ("highschool", r"\b(high school diploma|g\.?e\.?d)\b"),
]
DEGREE_ORDER = ["highschool", "associate", "bachelors", "masters", "doctorate"]

_SECTIONS = {
    "skills": r"^\s*(technical\s+)?skills?(\s*&\s*\w+)?\s*:?\s*$",
    "education": r"^\s*education\s*:?\s*$",
    "experience": r"^\s*(work\s+|professional\s+|relevant\s+)?experience\s*:?\s*$",
    "projects": r"^\s*(personal\s+|technical\s+)?projects?\s*:?\s*$",
}

_DATE_RANGE = re.compile(
    r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4}|\d{1,2}/\d{4}|\d{4})"
    r"\s*(?:-|–|—|to)\s*"
    r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4}|\d{1,2}/\d{4}|\d{4}|present|current)",
    re.I,
)

_SENIORITY = [
    ("intern", r"\b(intern|internship|co-?op)\b"),
    ("new-grad", r"\b(new\s?grad|entry[- ]level|university grad|campus hire)\b"),
    ("junior", r"\b(junior|jr\.?|associate|i{1,2}\b|\b1\b)\b"),
    ("senior", r"\b(senior|sr\.?|staff|principal|lead)\b"),
]

_SINGLE_DATE = re.compile(
    r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4}|\d{1,2}/\d{4}|\b(?:19|20)\d{2}\b)",
    re.I,
)

_BULLET = re.compile(r"^\s*[•·▪◦*\-–—]\s+")


def extract_text(path: Path) -> str:
    """Pull plain text out of a resume file. Raw text is retained so parsing can be redone."""
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED:
        raise ValueError(
            f"unsupported resume format {suffix!r}; expected {sorted(SUPPORTED)}"
        )
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        import docx

        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    return path.read_text(encoding="utf-8", errors="replace")


def split_sections(text: str) -> dict[str, list[str]]:
    """Bucket lines under the resume's own section headers."""
    lines = text.splitlines()
    out: dict[str, list[str]] = {"_preamble": []}
    current = "_preamble"
    for line in lines:
        matched = None
        for name, pattern in _SECTIONS.items():
            if re.match(pattern, line, re.I):
                matched = name
                break
        if matched:
            current = matched
            out.setdefault(current, [])
            continue
        out.setdefault(current, []).append(line)
    return out


def parse_skills(section: list[str]) -> list[str]:
    """Skills lines are usually comma- or bullet-delimited, often with a 'Languages:' label."""
    skills: list[str] = []
    for line in section:
        line = _BULLET.sub("", line).strip()
        if not line:
            continue
        _, _, tail = line.partition(":") if ":" in line[:30] else ("", "", line)
        for token in re.split(r"[,;|/•]", tail or line):
            token = token.strip(" .\t")
            if 1 < len(token) <= 40 and not token.lower().startswith("http"):
                skills.append(token)
    seen: set[str] = set()
    out: list[str] = []
    for s in skills:
        key = s.lower()
        if key not in seen:
            seen.add(key)
            out.append(s)
    return out


def detect_degree_level(text: str) -> str | None:
    found = [level for level, pattern in _DEGREE_LEVELS if re.search(pattern, text, re.I)]
    if not found:
        return None
    return max(found, key=lambda lvl: DEGREE_ORDER.index(lvl))


def parse_education(section: list[str], full_text: str) -> list[Education]:
    entries: list[Education] = []
    for line in section:
        stripped = line.strip()
        if not stripped or len(stripped) < 4:
            continue
        level = detect_degree_level(stripped)
        looks_like_school = re.search(
            r"\b(university|college|institute|school|academy)\b", stripped, re.I
        )
        if not level and not looks_like_school:
            continue
        # A graduation date is usually a single date ("May 2026"), not a range. Fall back to
        # the last standalone date on the line when there is no range to read.
        span = _DATE_RANGE.search(stripped)
        if span:
            graduation = span.group(2)
        else:
            singles = _SINGLE_DATE.findall(stripped)
            graduation = singles[-1] if singles else None
        entries.append(
            Education(
                school=stripped if looks_like_school else None,
                degree=stripped if level else None,
                degree_level=level,
                graduation=graduation,
                source_span=stripped[:200],
            )
        )
    return entries


def parse_employment(section: list[str]) -> list[Employment]:
    """Employer/title/date triples. Conservative: a line without a date range is not a job."""
    out: list[Employment] = []
    for line in section:
        stripped = line.strip()
        if not stripped or _BULLET.match(line):
            continue
        m = _DATE_RANGE.search(stripped)
        if not m:
            continue
        head = stripped[: m.start()].strip(" ,|–—-\t")
        parts = re.split(r"\s*[|–—]\s*|\s{2,}|\s+[-]\s+", head)
        parts = [p.strip(" ,") for p in parts if p.strip(" ,")]
        title = parts[0] if parts else None
        employer = parts[1] if len(parts) > 1 else None
        out.append(
            Employment(
                employer=employer, title=title,
                start=m.group(1), end=m.group(2),
                source_span=stripped[:200],
            )
        )
    return out


def detect_seniority(text: str) -> str | None:
    for level, pattern in _SENIORITY:
        if re.search(pattern, text, re.I):
            return level
    return None


def parse_profile(texts: dict[str, str]) -> Profile:
    """Build a structured profile from one or more resume texts.

    Deterministic only. Anything this cannot segment is left empty for the user to fill in
    profile.yaml, or for an optional LLM pass to propose with a source span.
    """
    profile = Profile(parsed_from=sorted(texts))
    all_skills: list[str] = []
    for name, text in texts.items():
        sections = split_sections(text)
        all_skills.extend(parse_skills(sections.get("skills", [])))
        profile.education.extend(parse_education(sections.get("education", []), text))
        profile.employment.extend(parse_employment(sections.get("experience", [])))
        if profile.name is None:
            for line in text.splitlines()[:5]:
                candidate = line.strip()
                if 2 <= len(candidate.split()) <= 5 and re.fullmatch(
                    r"[A-Za-z.\-' ]{4,60}", candidate
                ):
                    profile.name = candidate
                    break
        del name

    seen: set[str] = set()
    for s in all_skills:
        if s.lower() not in seen:
            seen.add(s.lower())
            profile.skills.append(s)

    joined = "\n".join(texts.values())
    profile.highest_degree_level = detect_degree_level(joined)
    profile.seniority = detect_seniority(joined)
    return profile


#: Words too common in any resume to say anything about a document's emphasis. Three groups:
#: ordinary filler, resume-structural vocabulary (every resume has an "Education" section, so
#: the word carries no signal), and dates.
_STOP = {
    # filler
    "and", "the", "for", "with", "from", "that", "this", "using", "used", "our", "was", "were",
    "have", "has", "had", "will", "would", "into", "over", "under", "team", "work", "worked",
    "role", "years", "year", "including", "across", "within", "also", "new", "use", "developed",
    "built", "designed", "created", "implemented", "responsible", "via", "per", "such",
    # resume structure — present in every document, so it distinguishes none of them
    "education", "experience", "skills", "skill", "projects", "project", "summary", "objective",
    "activities", "awards", "honors", "certifications", "coursework", "relevant", "interests",
    "languages", "tools", "technologies", "technical", "professional", "employment", "history",
    "university", "college", "school", "institute", "bachelor", "bachelors", "master", "masters",
    "degree", "gpa", "minor", "major", "present", "current", "graduated", "expected",
    # dates
    "january", "february", "march", "april", "may", "june", "july", "august", "september",
    "october", "november", "december", "jan", "feb", "mar", "apr", "jun", "jul", "aug", "sep",
    "sept", "oct", "nov", "dec",
}


def extract_emphasis(
    text: str, limit: int = 25, *, exclude: set[str] | None = None
) -> tuple[list[str], str]:
    """What this document leads with.

    Term frequency, weighted toward the first third of the page — a resume puts what it wants
    read first at the top. This is how the tool picks *which* document to send; it never
    rewrites one.

    ``exclude`` drops terms that carry no signal for *this* document — the candidate's own name
    appears in all of their resumes, so it cannot distinguish between them.
    """
    words = re.findall(r"[A-Za-z][A-Za-z+#./-]{2,}", text)
    if not words:
        return [], ""
    skip = _STOP | (exclude or set())
    cutoff = max(1, len(words) // 3)
    scores: dict[str, float] = {}
    for i, raw in enumerate(words):
        word = raw.lower().strip("./-")
        if len(word) < 3 or word in skip:
            continue
        scores[word] = scores.get(word, 0.0) + (2.0 if i < cutoff else 1.0)
    ranked = sorted(scores.items(), key=lambda kv: (-kv[1], kv[0]))[:limit]
    terms = [t for t, _ in ranked]
    return terms, "leads with " + ", ".join(terms[:6]) if terms else ""


def add_resume(conn: sqlite3.Connection, path: Path, *, is_master: bool = False) -> ResumeDoc:
    """Register a resume the user maintains. The file itself stays where it is."""
    path = path.expanduser().resolve()
    if not path.exists():
        raise FileNotFoundError(path)
    text = extract_text(path)
    if not text.strip():
        raise ValueError(f"no text extracted from {path.name} (is it a scanned image?)")
    digest = hashlib.sha256(text.encode("utf-8")).hexdigest()
    # The candidate's own name is in every one of their resumes, so it distinguishes none.
    parsed_name = parse_profile({path.name: text}).name or ""
    terms, summary = extract_emphasis(text, exclude=set(parsed_name.lower().split()))
    conn.execute(
        """INSERT INTO resumes (filename, path, added_at, text_sha256, raw_text,
                                emphasis_terms, emphasis_summary, is_master)
           VALUES (?,?,?,?,?,?,?,?)
           ON CONFLICT(path) DO UPDATE SET
               text_sha256=excluded.text_sha256, raw_text=excluded.raw_text,
               emphasis_terms=excluded.emphasis_terms,
               emphasis_summary=excluded.emphasis_summary, is_master=excluded.is_master""",
        (path.name, str(path), iso(utcnow()), digest, text,
         jdump(terms), summary, int(is_master)),
    )
    row = conn.execute("SELECT * FROM resumes WHERE path = ?", (str(path),)).fetchone()
    return _row_to_doc(row)


def list_resumes(conn: sqlite3.Connection) -> list[ResumeDoc]:
    rows = conn.execute("SELECT * FROM resumes ORDER BY id")
    return [_row_to_doc(r) for r in rows]


def resume_texts(conn: sqlite3.Connection) -> dict[str, str]:
    rows = conn.execute("SELECT filename, raw_text FROM resumes")
    return {r["filename"]: r["raw_text"] for r in rows}


def _row_to_doc(row: sqlite3.Row) -> ResumeDoc:
    return ResumeDoc(
        id=row["id"], filename=row["filename"], path=row["path"],
        text_sha256=row["text_sha256"],
        emphasis_terms=jload(row["emphasis_terms"], []) or [],
        emphasis_summary=row["emphasis_summary"], is_master=bool(row["is_master"]),
    )
